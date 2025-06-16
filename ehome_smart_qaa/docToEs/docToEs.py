import os
import elasticsearch
import requests
import time
import PyPDF2
from docx import Document
import spacy

def pdf_to_string(file_path):
    text = ""
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text()
    return text

def pdf_content(word_file, nlp):
    pdf_text = pdf_to_string(word_file)
    
    doc = nlp(pdf_text)

    sentences = list(doc.sents)

    buffer = ''
    buffer_size = 200
    chunks =[]

    for s in sentences:
        if len(buffer) <= buffer_size:
            buffer = buffer + str(s)
        else:
            chunks.append(buffer)
            buffer = ''

    return chunks

def docx_content(word_file, nlp):
    doc = Document(word_file)

    full_text = []
    for para in doc.paragraphs:
        para_text = []
        for run in para.runs:
            if run._element.xpath('.//w:br'):
                para_text.append(run.text + '\n')
            else:
                para_text.append(run.text)
        full_text.append(''.join(para_text))

    full_text_str = ''
    for i in full_text:
        full_text_str += i
        # full_text_str += i + '\n'

    doc = nlp(full_text_str)

    sentences = list(doc.sents)

    buffer = ''
    buffer_size = 200
    chunks =[]

    for s in sentences:
        if len(buffer) <= buffer_size:
            buffer = buffer + str(s)
        else:
            chunks.append(buffer)
            buffer = ''
            
    chunks.append(buffer)

    return chunks

def embeddingMsg(sentence):
    # api_key = "sk-JyubfGLXR4nbxkL3C801Cc1c61F743749e40Ac1688Af870c"
    api_key = "sk-1cRpS3i6xqjeT3bCE018D16aB0274124B2Da8e0f6e85E5F7"
    headers = {
        "Authorization": f"Bearer {api_key}",  # API key in the Authorization header
        "Content-Type": "application/json"  # Specify the content type
    }

    # url = 'http://192.168.24.137:30123/open-api/oneapi/v1/embeddings'
    url = 'http://172.22.80.93:38080/open-api/2/oneapi/v1/embeddings'
    myobj = {"input":[sentence],
        "model":"text-embedding-3-large"}
    
    x = requests.post(url, json = myobj, headers = headers)
    res = x.json()["data"][0]["embedding"]
    
    return res

def chunks_to_es(chunks,word_file,counter, ESindex):
    es = elasticsearch.Elasticsearch(
         hosts = ['http://172.17.172.149:9200'],
         basic_auth = ('es2_zntj', 'Zzty240603!@#')
        #   hosts=['http://192.168.152.47:9292'],
        #   basic_auth=('ai', 'Aireccontent@123')
        )
    
    for c in chunks:
        insertdoc = {}
        
        cur = time.localtime()
        id = 'fa'+str(counter)

        
        insertdoc["searhShow"] = "1"
        insertdoc["question"] = c
        insertdoc["modelId"] = "1"
        insertdoc["secondLevelDirectory"] = word_file
        insertdoc["extendIssues"] = []
        insertdoc["updateTime"] =str(time.asctime(cur))
        insertdoc["showScope"] = "2"
        insertdoc["hot"] = 0
        insertdoc["thirdLevelDirectory"] = word_file
        insertdoc["answer"] = c
        insertdoc["auditStatus"] =  "2"
        insertdoc["id"] = "1"
        insertdoc["keyword"] = "1"
        insertdoc["firstLevelDirectory"] = "计财部门"
        insertdoc["question_embedding"] =embeddingMsg(c)
        
        res = es.index(index=ESindex, id=counter, document=insertdoc)
        print(res['result'])  # 应该返回 'created' 或 'updated'
        
        counter += 1
    
    fres = res['result'] + ',counter number is: ' + str(counter)
        
    return fres
        
def ESupload(file, ESindex, counter):
    word_file = file
    
    nlp = spacy.load("zh_core_web_sm")
    
    chunks = docx_content(word_file, nlp)
    
    upload_res = chunks_to_es(chunks,word_file,counter,ESindex)
    
    return upload_res

def find_table_line_numbers(doc_path):
    doc = Document(doc_path)
    results = []
    
    # 获取文档的XML结构
    doc_element = doc.element.body
    
    # 收集所有段落和表格的位置信息
    elements = []
    line_counter = 1
    
    for child in doc_element.iterchildren():
        if child.tag.endswith('p'):  # 段落
            num_lines = len(child.text.split('\n')) if child.text else 1
            elements.append({
                'type': 'paragraph',
                'start_line': line_counter,
                'end_line': line_counter + num_lines - 1,
                'element': child
            })
            line_counter += num_lines
        elif child.tag.endswith('tbl'):  # 表格
            # 估算表格行数(每行单元格数)
            table = Table(child, doc)
            num_lines = sum(1 for _ in table.rows)
            elements.append({
                'type': 'table',
                'start_line': line_counter,
                'end_line': line_counter + num_lines - 1,
                'element': child,
                'table_obj': table
            })
            line_counter += num_lines
    
    # 提取表格位置信息
    for elem in elements:
        if elem['type'] == 'table':
            # 查找前一个非空段落
            preceding_text = None
            for prev_elem in reversed(elements[:elements.index(elem)]):
                if prev_elem['type'] == 'paragraph' and prev_elem['element'].text.strip():
                    preceding_text = prev_elem['element'].text.strip()
                    break
            
            results.append({
                'table_index': len(results) + 1,
                'start_line': elem['start_line'],
                'end_line': elem['end_line'],
                'preceding_text': preceding_text,
                'table': elem['table_obj']
            })
    
    return results

def list_to_markdown_table(data):
    # 转义管道符
    escape_pipes = lambda s: str(s).replace("|", "\|") if s else ""
    # 构建表格
    header = "| " + " | ".join(escape_pipes(cell) for cell in data[0]) + " |\n"
    separator = "| " + " | ".join(["---"] * len(data[0])) + " |\n"
    body = ""
    for row in data[1:]:
        body += "| " + " | ".join(escape_pipes(cell) for cell in row) + " |\n"
    return header + separator + body

def docTableUpload(word_file, esindex, counter):

    es = elasticsearch.Elasticsearch(
        #  hosts = ['http://172.17.172.149:9200'],
        #  basic_auth = ('es2_zntj', 'Zzty240603!@#')
          hosts=['http://192.168.152.47:9292'],
          basic_auth=('ai', 'Aireccontent@123')
        )

    table_info = find_table_line_numbers(word_file)

    for info in table_info:
        table_content = []
        for row in info['table'].rows:
            row_content = [cell.text.strip() for cell in row.cells]
            table_content.append(row_content)
        # print(table_content)
        insertv = list_to_markdown_table(table_content)

        insertdoc = {}

        cur = time.localtime()
        id = 'fa'+str(counter)

        insertdoc["searhShow"] = "1"
        insertdoc["question"] = info['preceding_text']
        insertdoc["modelId"] = "1"
        insertdoc["secondLevelDirectory"] = word_file
        insertdoc["extendIssues"] = []
        insertdoc["updateTime"] =str(time.asctime(cur))
        insertdoc["showScope"] = "2"
        insertdoc["hot"] = 0
        insertdoc["thirdLevelDirectory"] = word_file
        insertdoc["answer"] = insertv
        insertdoc["auditStatus"] =  "2"
        insertdoc["id"] = id
        insertdoc["keyword"] = "1"
        insertdoc["firstLevelDirectory"] = "计财部门"
        insertdoc["question_embedding"] =embeddingMsg(insertv)

        res = es.index(index=esindex, id=counter, document=insertdoc)
        print(res['result'])  # 应该返回 'created' 或 'updated'
        
        # print(type(insertv))
        # print(insertdoc)
        counter += 1
        print(id)
        print('#'*40)

    return counter
# counter = 289

# chunks = docx_content(word_file, nlp)

# chunks_to_es(chunks,word_file,counter)

# word_file = '关于印发《申万宏源集团股份有限公司和申万宏源证券有限公司异地工作人员福利待遇管理办法（2024年修订）》的通知（申万宏源党发〔2024〕116号）.pdf'

# chunks = pdf_content(word_file, nlp)

# chunks_to_es(chunks,word_file,counter)

# testq = '怎么用公务卡进行报销'
# query = {
#   "query": {
#     "script_score": {
#       "query": {"match_all": {}},
#       "script": {
#         "source": """
#           cosineSimilarity(params.embtestq, 'question_embedding') + 1.0
#         """,
#         "params": {"embtestq": embtestq}
#       }
#     }
#   },
#   "size": 3
# }
# results = es.search(index="ehome_smart_qaa_fa", body=query)

# folder_path = '''C:/Users/Luyang/Documents/account_kms/files'''
# all_items = os.listdir(folder_path)
# files = [item for item in all_items if os.path.isfile(os.path.join(folder_path, item))]
# for file in files:
#     if file[-5:] != 'ipynb':
#         print(file)